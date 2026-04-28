# Tarsim API v0.5
# AI Resume Tailoring Tool

import os
from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import pdfplumber
from openai import OpenAI
import io
import requests as http_requests
from docx import Document
from docx.shared import Pt, Inches

app = FastAPI(title='Tarsim API', version='0.5')

app.add_middleware(
    CORSMiddleware,
    allow_origins=['*'],
    allow_methods=['*'],
    allow_headers=['*'],
)

# OpenAI setup
OPENAI_KEY = os.environ.get('OPENAI_KEY', '')
openai_client = OpenAI(api_key=OPENAI_KEY)

# SerpAPI setup
SERPAPI_KEY = os.environ.get('SERPAPI_KEY', '')


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract all text from a PDF file."""
    text = ''
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'
    return text.strip()


def analyze_resume(resume_text: str) -> dict:
    """Send resume text to OpenAI for analysis."""
    prompt = f"""Analyze this resume and provide constructive improvement suggestions.

Resume text:
{resume_text}

Provide your response in this exact format:
1. STRENGTHS (3 specific strengths)
2. WEAKNESSES (3 specific weaknesses)
3. ATS_OPTIMIZATION (3 specific ATS keyword/formatting suggestions)
4. ACTION_VERB_UPGRADES (3 weak verbs to replace with stronger ones, format: "weak -> strong")
5. OVERALL_SCORE (out of 10)

Be specific. Reference exact text from the resume."""

    response = openai_client.chat.completions.create(
        model='gpt-4o-mini',
        messages=[{'role': 'user', 'content': prompt}],
        max_tokens=800,
        temperature=0.7
    )

    return {
        'analysis': response.choices[0].message.content,
        'tokens_used': response.usage.total_tokens
    }


def tailor_resume(resume_text: str, job_description: str) -> dict:
    """Tailor resume to a specific job description and generate a cover letter."""
    prompt = f"""You are an expert resume writer. Tailor the candidate's resume to match this specific job description, and write a custom cover letter.

CANDIDATE'S RESUME:
{resume_text}

JOB DESCRIPTION:
{job_description}

Provide your response in this exact format:

===TAILORED_RESUME===
(Rewrite the resume to emphasize skills/experience matching this job. Keep all factual content true to the original — don't invent experience. Reorder bullets, rephrase descriptions, and integrate relevant keywords from the job description naturally. Output the full tailored resume as plain text.)

===COVER_LETTER===
(Write a 3-paragraph cover letter for this specific role. Tone: confident, professional, not generic. Reference specific points from the job description. Show genuine fit, not desperation.)

===KEY_CHANGES===
(Bullet point list of the main changes you made and why.)

===MATCH_SCORE===
(Score out of 10 for how well the candidate fits this role based on their existing experience.)

Be specific. Don't fabricate. Make it sound human, not robotic."""

    response = openai_client.chat.completions.create(
        model='gpt-4o-mini',
        messages=[{'role': 'user', 'content': prompt}],
        max_tokens=2000,
        temperature=0.7
    )

    return {
        'output': response.choices[0].message.content,
        'tokens_used': response.usage.total_tokens
    }


def search_jobs(query: str, location: str = '', num_results: int = 20) -> dict:
    """Search jobs via SerpAPI Google Search engine using boolean operators."""

    boolean_query = f'({query}) (site:linkedin.com/jobs OR site:indeed.com/viewjob OR site:glassdoor.com/job-listing OR site:ziprecruiter.com/jobs)'

    if location:
        boolean_query += f' "{location}"'

    url = 'https://serpapi.com/search'
    params = {
        'engine': 'google',
        'q': boolean_query,
        'api_key': SERPAPI_KEY,
        'num': num_results,
    }

    response = http_requests.get(url, params=params, timeout=30)
    response.raise_for_status()
    data = response.json()

    jobs = []
    for result in data.get('organic_results', []):
        jobs.append({
            'title': result.get('title', ''),
            'link': result.get('link', ''),
            'snippet': result.get('snippet', ''),
            'source': result.get('displayed_link', '').split('/')[0] if result.get('displayed_link') else '',
        })

    return {
        'query': boolean_query,
        'total_results': len(jobs),
        'jobs': jobs
    }


def text_to_docx(text: str) -> io.BytesIO:
    """Convert plain text resume/cover letter to a Word document."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Add content paragraph by paragraph
    for line in text.split('\n'):
        line = line.strip()
        if line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


@app.get('/')
def home():
    return {
        'name': 'Tarsim API',
        'version': '0.5',
        'description': 'AI Resume Tailoring Tool',
        'status': 'running',
        'endpoints': [
            '/',
            '/analyze-resume',
            '/tailor-resume',
            '/search-jobs',
            '/export-docx'
        ]
    }


@app.post('/analyze-resume')
async def analyze_resume_endpoint(file: UploadFile = File(...)):
    """Upload a PDF resume and get improvement suggestions."""

    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail='File must be a PDF')

    try:
        pdf_bytes = await file.read()
        resume_text = extract_text_from_pdf(pdf_bytes)

        if not resume_text:
            raise HTTPException(status_code=400, detail='No text found in PDF')

        result = analyze_resume(resume_text)

        return {
            'filename': file.filename,
            'text_length': len(resume_text),
            'analysis': result['analysis'],
            'tokens_used': result['tokens_used']
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post('/tailor-resume')
async def tailor_resume_endpoint(
    file: UploadFile = File(...),
    job_description: str = Form(...)
):
    """Upload a PDF resume + job description, get tailored resume and cover letter."""

    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail='File must be a PDF')

    if len(job_description.strip()) < 50:
        raise HTTPException(status_code=400, detail='Job description too short (min 50 chars)')

    try:
        pdf_bytes = await file.read()
        resume_text = extract_text_from_pdf(pdf_bytes)

        if not resume_text:
            raise HTTPException(status_code=400, detail='No text found in PDF')

        result = tailor_resume(resume_text, job_description)

        return {
            'filename': file.filename,
            'resume_length': len(resume_text),
            'job_description_length': len(job_description),
            'output': result['output'],
            'tokens_used': result['tokens_used']
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get('/search-jobs')
def search_jobs_endpoint(query: str, location: str = ''):
    """Search jobs across major job boards using boolean search via Google."""

    if not SERPAPI_KEY:
        raise HTTPException(status_code=500, detail='SerpAPI key not configured')

    if len(query.strip()) < 3:
        raise HTTPException(status_code=400, detail='Query too short (min 3 chars)')

    try:
        result = search_jobs(query, location)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post('/export-docx')
async def export_docx_endpoint(
    text: str = Form(...),
    filename: str = Form('document')
):
    """Convert text to a downloadable Word document."""

    if len(text.strip()) < 10:
        raise HTTPException(status_code=400, detail='Text too short')

    try:
        docx_file = text_to_docx(text)

        return StreamingResponse(
            docx_file,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="{filename}.docx"'}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == '__main__':
    import uvicorn
    print('=== Tarsim API v0.5 ===')
    print('Starting server on http://localhost:8000')
    uvicorn.run(app, host='0.0.0.0', port=8000)